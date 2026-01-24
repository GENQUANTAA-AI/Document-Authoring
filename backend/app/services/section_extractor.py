"""
Service for extracting sections from DOCX templates
Supports:
1. Word Heading styles (Heading 1, Heading 2)
2. ALL CAPS text ending with ":" (e.g. PURPOSE:, PROCEDURE:)
"""
from docx import Document
from docx.shared import Pt
from typing import List, Tuple, Optional
import re

class SectionExtractor:
    """Extracts sections from DOCX documents"""
    
    @staticmethod
    def extract_sections(docx_path: str) -> List[Tuple[str, int]]:
        """
        Extract sections from DOCX file.
        Returns list of tuples: (section_title, paragraph_index)
        """
        doc = Document(docx_path)
        sections = []
        
        # Pattern for ALL CAPS text ending with ":"
        all_caps_pattern = re.compile(r'^[A-Z][A-Z\s]+:$')
        
        for i, para in enumerate(doc.paragraphs):
            # Check for Heading styles
            if para.style.name.startswith('Heading'):
                # Extract heading level (1, 2, etc.)
                heading_level = int(para.style.name.replace('Heading ', ''))
                if heading_level <= 2:  # Only Heading 1 and Heading 2
                    text = para.text.strip()
                    if text:
                        sections.append((text, i))
            
            # Check for ALL CAPS pattern
            else:
                text = para.text.strip()
                if all_caps_pattern.match(text):
                    sections.append((text, i))
        
        # Remove duplicates while preserving order
        seen = set()
        unique_sections = []
        for title, idx in sections:
            if title not in seen:
                seen.add(title)
                unique_sections.append((title, idx))
        
        return unique_sections
    
    @staticmethod
    def get_section_range(doc: Document, section_start_idx: int, next_section_idx: Optional[int] = None) -> List:
        """
        Get all paragraphs for a section (from section_start_idx to next_section_idx)
        """
        if next_section_idx is None:
            return doc.paragraphs[section_start_idx:]
        return doc.paragraphs[section_start_idx:next_section_idx]

