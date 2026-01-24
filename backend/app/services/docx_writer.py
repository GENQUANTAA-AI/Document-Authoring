"""
Service for writing generated content to DOCX documents
"""
from docx import Document
from docx.table import Table
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import List, Optional, Tuple
import re
from app.services.sop_numbering import SOPNumbering

class DOCXWriter:
    """Handles writing content to DOCX documents"""
    
    @staticmethod
    def find_section_by_title(doc: Document, section_title: str, start_search_idx: int = 0) -> Optional[int]:
        """
        Find section header by matching title.
        Returns paragraph index of matching section header.
        """
        normalized_title = section_title.strip().upper()
        all_caps_pattern = re.compile(r'^[A-Z][A-Z\s]+:$')
        
        for i in range(start_search_idx, len(doc.paragraphs)):
            para = doc.paragraphs[i]
            para_text = para.text.strip()
            
            # Check if it's a heading style
            if para.style.name.startswith('Heading'):
                heading_level = int(para.style.name.replace('Heading ', ''))
                if heading_level <= 2:
                    # Compare normalized text
                    if para_text.upper() == normalized_title or para_text.upper() == normalized_title.replace(':', ''):
                        return i
            
            # Check if it matches ALL CAPS pattern
            if all_caps_pattern.match(para_text):
                if para_text.upper() == normalized_title or para_text.upper() == normalized_title.replace(':', ''):
                    return i
        
        return None
    
    @staticmethod
    def has_table_in_section(doc: Document, section_start_idx: int, next_section_idx: Optional[int]) -> Tuple[bool, Optional[Table]]:
        """
        Check if section contains a table.
        Returns (has_table, table_object)
        """
        # Get document body element
        body = doc.element.body
        
        # Find section header paragraph element
        section_para = doc.paragraphs[section_start_idx]
        section_para_element = section_para._element
        
        # Get all elements in document body
        all_elements = list(body)
        section_element_idx = all_elements.index(section_para_element)
        
        # Determine end index
        if next_section_idx is not None and next_section_idx < len(doc.paragraphs):
            next_section_para = doc.paragraphs[next_section_idx]
            next_section_element = next_section_para._element
            try:
                end_idx = all_elements.index(next_section_element)
            except ValueError:
                end_idx = len(all_elements)
        else:
            end_idx = len(all_elements)
        
        # Check for tables between section header and next section
        for i in range(section_element_idx + 1, end_idx):
            element = all_elements[i]
            if element.tag.endswith('tbl'):  # Table element
                # Find the corresponding table object
                for table in doc.tables:
                    if table._element == element:
                        return True, table
                # If table found but not in doc.tables, create a reference
                return True, None
        
        return False, None
    
    @staticmethod
    def fill_table_content(table, content: str, section_number: int):
        """
        Fill content into a table.
        For annexures, formats, reference documents, etc.
        Handles different table structures intelligently.
        """
        # Apply SOP numbering
        numbered_content = SOPNumbering.apply_numbering(content, section_number)
        
        # Split content into lines (handle both paragraphs and bullet points)
        lines = []
        for line in numbered_content.split('\n'):
            stripped = line.strip()
            if stripped:
                # If it's a numbered item (N.N or N.N.N), use it as-is
                if re.match(r'^\d+\.\d+', stripped):
                    lines.append(stripped)
                # Otherwise, treat as paragraph content
                elif not stripped.startswith('-') and not stripped.startswith('--'):
                    lines.append(stripped)
        
        if not lines:
            return
        
        # Determine if table has headers (check if first row has non-empty cells)
        has_headers = False
        if len(table.rows) > 0:
            first_row_cells = [cell.text.strip() for cell in table.rows[0].cells]
            has_headers = any(cell and not cell.startswith(('1.', '2.', '3.', '4.', '5.')) for cell in first_row_cells)
        
        # Start from row 1 if headers exist, else row 0
        start_row = 1 if has_headers else 0
        
        # Determine number of columns
        num_cols = len(table.rows[0].cells) if table.rows else 1
        
        # Fill content into table rows
        for i, line in enumerate(lines):
            row_idx = start_row + i
            
            # Add new row if needed
            if row_idx >= len(table.rows):
                new_row = table.add_row()
                # Ensure new row has same number of columns
                while len(new_row.cells) < num_cols:
                    new_row.add_cell()
            
            # Fill content into cells
            # If single column, fill first cell
            # If multiple columns, try to parse and distribute
            if num_cols == 1:
                table.rows[row_idx].cells[0].text = line
            else:
                # For multi-column tables, try to split by common delimiters
                # Otherwise, put everything in first column
                parts = re.split(r'\s*[|,;]\s*', line, maxsplit=num_cols - 1)
                for col_idx, part in enumerate(parts[:num_cols]):
                    if col_idx < len(table.rows[row_idx].cells):
                        table.rows[row_idx].cells[col_idx].text = part.strip()
    
    @staticmethod
    def insert_section_content(
        doc: Document,
        section_start_idx: int,
        next_section_idx: Optional[int],
        content: str,
        section_number: int,
        section_title: str
    ):
        """
        Insert generated content into DOCX after section header.
        Validates section title before inserting to ensure content goes to correct section.
        
        Args:
            doc: Document object
            section_start_idx: Initial index of section header paragraph (may have shifted)
            next_section_idx: Index of next section header (None if last section)
            content: Generated content to insert
            section_number: Section number for SOP numbering
            section_title: Title of the section (for validation)
        """
        # First, validate and find the correct section header by matching title
        # This handles cases where indices have shifted due to previous insertions
        actual_section_idx = DOCXWriter.find_section_by_title(doc, section_title, max(0, section_start_idx - 5))
        
        if actual_section_idx is None:
            # Fallback to original index if title matching fails
            actual_section_idx = section_start_idx
        else:
            # Verify it's the right section by checking it matches expected position
            # Allow some flexibility (±5 paragraphs) for minor shifts
            if abs(actual_section_idx - section_start_idx) > 10:
                # If found section is too far from expected, use original
                actual_section_idx = section_start_idx
        
        # Validate the section header text matches
        section_para = doc.paragraphs[actual_section_idx]
        section_para_text = section_para.text.strip().upper()
        expected_title = section_title.strip().upper()
        
        # Check if this is actually the section header we want
        is_valid_section = (
            section_para.style.name.startswith('Heading') or
            re.match(r'^[A-Z][A-Z\s]+:$', section_para_text)
        ) and (
            section_para_text == expected_title or
            section_para_text == expected_title.replace(':', '') or
            expected_title in section_para_text or
            section_para_text in expected_title
        )
        
        if not is_valid_section:
            raise ValueError(
                f"Section header mismatch! Expected '{section_title}' at index {actual_section_idx}, "
                f"but found '{section_para.text.strip()}'"
            )
        
        # Check if section has a table (for annexures, formats, etc.)
        has_table, table = DOCXWriter.has_table_in_section(doc, actual_section_idx, next_section_idx)
        
        if has_table and table:
            # Fill content into table
            DOCXWriter.fill_table_content(table, content, section_number)
            return
        
        # Otherwise, insert as paragraphs
        # Apply SOP numbering
        numbered_content = SOPNumbering.apply_numbering(content, section_number)
        
        # Split content into elements
        elements = SOPNumbering.split_into_paragraphs_and_bullets(numbered_content)
        
        # Get the section header paragraph and its parent element
        section_para_element = section_para._element
        parent = section_para_element.getparent()
        
        # Get the index where we need to insert (after section header)
        insert_index = parent.index(section_para_element) + 1
        
        # Build all paragraphs to insert
        paragraphs_to_insert = []
        
        for element in elements:
            if element['type'] == 'paragraph':
                # Create paragraph with content
                para = doc.add_paragraph(element['content'])
                para.style = 'Normal'
                paragraphs_to_insert.append(para._element)
            
            elif element['type'] == 'bullet_list':
                # Add bullet list items
                for item in element['items']:
                    item_text = item
                    
                    # Check if it's a sub-bullet (N.N.N pattern)
                    if re.match(r'^\d+\.\d+\.\d+\s+', item_text):
                        # Sub-bullet - keep full text with numbering
                        para = doc.add_paragraph(item_text)
                        try:
                            para.style = 'List Bullet 2'
                        except KeyError:
                            para.style = 'Normal'
                        para.paragraph_format.left_indent = Inches(0.5)
                        para.paragraph_format.first_line_indent = Inches(-0.25)
                    # Check if it's a main bullet (N.N pattern)
                    elif re.match(r'^\d+\.\d+\s+', item_text):
                        # Main bullet - keep full text with numbering
                        para = doc.add_paragraph(item_text)
                        try:
                            para.style = 'List Bullet'
                        except KeyError:
                            para.style = 'Normal'
                        para.paragraph_format.left_indent = Inches(0.25)
                        para.paragraph_format.first_line_indent = Inches(-0.25)
                    else:
                        # Fallback - treat as regular paragraph
                        para = doc.add_paragraph(item_text)
                        para.style = 'Normal'
                    
                    paragraphs_to_insert.append(para._element)
        
        # Insert all paragraphs at the correct position (in reverse to maintain order)
        for para_element in reversed(paragraphs_to_insert):
            parent.remove(para_element)
            parent.insert(insert_index, para_element)
    
    @staticmethod
    def create_final_document(
        template_path: str,
        sections: List[dict],
        section_indices: dict,
        output_path: str
    ):
        """
        Create final document with all generated content.
        Validates each section before inserting content.
        
        Args:
            template_path: Path to original template
            sections: List of sections with content {number, title, content}
            section_indices: Dict mapping section number to paragraph index
            output_path: Path to save final document
        """
        doc = Document(template_path)
        
        # Process each section in order
        section_numbers = sorted(section_indices.keys())
        
        for i, section_num in enumerate(section_numbers):
            section = next((s for s in sections if s['number'] == section_num), None)
            if not section or not section.get('content'):
                continue
            
            section_start_idx = section_indices[section_num]
            section_title = section.get('title', '')
            next_section_idx = section_indices.get(section_numbers[i + 1]) if i + 1 < len(section_numbers) else None
            
            try:
                DOCXWriter.insert_section_content(
                    doc,
                    section_start_idx,
                    next_section_idx,
                    section['content'],
                    section_num,
                    section_title
                )
            except ValueError as e:
                # Log error but continue with other sections
                print(f"Warning: Failed to insert content for section {section_num} ({section_title}): {e}")
                continue
        
        # Save document
        doc.save(output_path)

